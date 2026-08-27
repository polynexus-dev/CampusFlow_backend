"""
NAAC SSR / AQAR document generator.

Compiles NAAC-tagged accreditation criteria (applied narrative(s) + evidence,
same source data SSRExportView already exposes as JSON+ZIP in
views/compliance.py) into a single submission-shaped .docx. Institution-wide,
not department/program-scoped — NAAC evaluates the institution as a whole,
unlike the NBA SAR (services/nba_sar_export.py), which is per-program.

Two modes, matching what an IQAC actually produces:
  - AQAR (financial_year given): one year's applied narrative + evidence per
    criterion — the annual incremental report.
  - SSR (financial_year omitted): every applied narrative across all years
    for each criterion, each labeled by year, plus the full multi-year
    evidence list. This is a deliberate improvement over SSRExportView's
    JSON, which drops narratives entirely when no year is given (narratives
    are year-scoped, so it only attaches one when the export itself is
    year-scoped) — for an actual SSR document that's too little to compile
    from, so here every already-applied narrative is surfaced instead,
    labeled by year, for IQAC to assemble the final SSR prose from. Nothing
    new is generated; only already-reviewed content is compiled.

Never generates new "official" content — a criterion/year with no applied
narrative renders an explicit placeholder, same boundary
AccreditationNarrativeDraft's own docstring describes.
"""
import io
from datetime import date

from django.db import connection

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..models.aqar_ssr import (
    AccreditationSubmission, FacultyResearchOutput, InstitutionalEvent, StudentFeedback,
)
from ..models.compliance import AccreditationCriterion, ComplianceCertificate, ComplianceCertificateType
from ..models.department import Department
from ..models.finance import FinancialYear
from ..models.academics import Program
from ..models.profile import (
    AdministratorProfile, DepartmentHeadProfile, ManagementProfile,
    NonTeachingStaffProfile, StudentProfile, TeachingStaffProfile,
)

AUDITED_FINANCIAL_STATEMENT_TYPE_NAME = "Audited Financial Statement"
TRAILING_FINANCIAL_YEARS = 5

NOT_REVIEWED_PLACEHOLDER = (
    "No IQAC-approved narrative has been applied for this criterion/year yet. "
    "Add evidence and request/apply an AI narrative draft before final submission."
)

NAAC_TOP_LEVEL_TITLES = {
    "1": "Curricular Aspects",
    "2": "Teaching-Learning and Evaluation",
    "3": "Research, Innovations and Extension",
    "4": "Infrastructure and Learning Resources",
    "5": "Student Support and Progression",
    "6": "Governance, Leadership and Management",
    "7": "Institutional Values and Best Practices",
}


def _add_table(document, header, rows):
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = "" if value is None else str(value)
    return table


def _title_page(document, financial_year):
    tenant_name = getattr(connection.tenant, "name", "") or connection.schema_name

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(tenant_name)
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if financial_year:
        run = subtitle.add_run(f"NAAC AQAR — Annual Quality Assurance Report ({financial_year.label})")
    else:
        run = subtitle.add_run("NAAC SSR — Self Study Report (all years, cumulative)")
    run.bold = True
    run.font.size = Pt(16)

    generated_line = document.add_paragraph()
    generated_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_line.add_run(f"Generated on {date.today().isoformat()} via CampusNexus")

    document.add_page_break()


def _extended_profile_section(document):
    document.add_heading("Institution Extended Profile", level=1)

    programs_by_level = list(
        Program.objects.filter(is_active=True).values_list("level", flat=True)
    )
    level_counts = {}
    for level in programs_by_level:
        level_counts[level] = level_counts.get(level, 0) + 1

    _add_table(
        document,
        ["Metric", "Value"],
        [
            ["Departments", Department.objects.count()],
            ["Active Programs", Program.objects.filter(is_active=True).count()],
            ["Total Students", StudentProfile.objects.count()],
            ["Students with Disability", StudentProfile.objects.filter(disability_status=True).count()],
            ["Teaching Staff", TeachingStaffProfile.objects.count()],
            ["Non-Teaching Staff", NonTeachingStaffProfile.objects.count()],
            ["Management Staff", ManagementProfile.objects.count()],
            ["Administrators", AdministratorProfile.objects.count()],
            ["Department Heads", DepartmentHeadProfile.objects.count()],
        ] + [[f"Programs — {level}", count] for level, count in sorted(level_counts.items())],
    )
    document.add_paragraph()


def _narrative_paragraphs_aqar(document, criterion, financial_year):
    narrative = criterion.narrative_drafts.filter(
        financial_year=financial_year, status="applied",
    ).order_by("-applied_at").first()
    document.add_paragraph(narrative.narrative_text if narrative else NOT_REVIEWED_PLACEHOLDER)


def _narrative_paragraphs_ssr(document, criterion):
    narratives = criterion.narrative_drafts.filter(status="applied").select_related(
        "financial_year",
    ).order_by("financial_year__start_date")
    if not narratives:
        document.add_paragraph(NOT_REVIEWED_PLACEHOLDER)
        return
    for narrative in narratives:
        p = document.add_paragraph()
        p.add_run(f"[{narrative.financial_year.label}] ").bold = True
        p.add_run(narrative.narrative_text)


def _criteria_section(document, financial_year):
    document.add_heading("Criterion-wise Narrative and Evidence", level=1)

    criteria = AccreditationCriterion.objects.filter(
        body=AccreditationCriterion.BODY_NAAC,
    ).order_by("code")

    current_group = None
    for criterion in criteria:
        group = criterion.code.split(".")[0]
        if group != current_group:
            current_group = group
            document.add_heading(
                f"Criterion {group}: {NAAC_TOP_LEVEL_TITLES.get(group, '')}", level=1,
            )

        document.add_heading(f"{criterion.code} — {criterion.title}", level=2)

        if financial_year:
            _narrative_paragraphs_aqar(document, criterion, financial_year)
        else:
            _narrative_paragraphs_ssr(document, criterion)

        evidence_qs = criterion.evidence.all()
        if financial_year:
            evidence_qs = evidence_qs.filter(financial_year=financial_year)
        evidence_items = list(evidence_qs.select_related("department", "financial_year", "linked_certificate").order_by("-created_at"))

        if not evidence_items:
            document.add_paragraph("No evidence recorded for this criterion/year yet.", style="Intense Quote")
            document.add_paragraph()
            continue

        header = ["Description", "Department", "Status"]
        if not financial_year:
            header.insert(1, "Year")

        rows = []
        for item in evidence_items:
            desc = item.description or (item.file.name.split("/")[-1] if item.file else "—")
            dept = item.department.name if item.department else "Institution-wide"
            row = [desc, dept, item.get_status_display()]
            if not financial_year:
                row.insert(1, item.financial_year.label)
            rows.append(row)

        _add_table(document, header, rows)
        document.add_paragraph()


def _faculty_output_section(document, financial_year):
    document.add_heading("Faculty Research Output", level=1)

    outputs = FacultyResearchOutput.objects.select_related("faculty__user", "financial_year")
    if financial_year:
        outputs = outputs.filter(financial_year=financial_year)
    outputs = list(outputs.order_by("faculty__employee_id", "-financial_year__start_date"))

    if not outputs:
        document.add_paragraph("No faculty research output recorded for this criterion/year yet.", style="Intense Quote")
        document.add_paragraph()
        return

    def _details(o):
        if o.output_type == FacultyResearchOutput.TYPE_PUBLICATION:
            return f"{o.journal_or_venue or '—'}" + (" (peer-reviewed)" if o.is_peer_reviewed else "")
        if o.output_type == FacultyResearchOutput.TYPE_GRANT:
            return f"{o.funding_agency or '—'} — Rs. {o.grant_amount_lakhs} lakhs"
        return f"{o.patent_number or '—'} ({o.get_patent_status_display() or 'status unspecified'})"

    header = ["Faculty", "Type", "Title", "Details"]
    if not financial_year:
        header.append("Year")
    rows = []
    for o in outputs:
        row = [o.faculty.user.get_full_name() or o.faculty.user.username, o.get_output_type_display(), o.title, _details(o)]
        if not financial_year:
            row.append(o.financial_year.label)
        rows.append(row)
    _add_table(document, header, rows)
    document.add_paragraph()


def _student_feedback_section(document, financial_year):
    document.add_heading("Student Feedback and Action Taken", level=1)

    feedback_qs = StudentFeedback.objects.select_related("department", "financial_year")
    if financial_year:
        feedback_qs = feedback_qs.filter(financial_year=financial_year)
    feedback_items = list(feedback_qs.order_by("-filed_date"))

    if not feedback_items:
        document.add_paragraph("No student feedback recorded for this criterion/year yet.", style="Intense Quote")
        document.add_paragraph()
        return

    header = ["Department", "Category", "Status", "Action Taken"]
    if not financial_year:
        header.insert(0, "Year")
    rows = []
    for f in feedback_items:
        row = [f.department.name if f.department else "Institution-wide", f.category or "—",
               f.get_status_display(), f.action_taken or "—"]
        if not financial_year:
            row.insert(0, f.financial_year.label)
        rows.append(row)
    _add_table(document, header, rows)
    document.add_paragraph()


def _institutional_events_section(document, financial_year):
    document.add_heading("Institutional Events and Activities", level=1)

    events_qs = InstitutionalEvent.objects.select_related("department", "financial_year")
    if financial_year:
        events_qs = events_qs.filter(financial_year=financial_year)
    events = list(events_qs.order_by("-event_date"))

    if not events:
        document.add_paragraph("No institutional events logged for this criterion/year yet.", style="Intense Quote")
        document.add_paragraph()
        return

    _add_table(
        document,
        ["Title", "Type", "Department", "Date", "Participants"],
        [
            [e.title, e.event_type or "—", e.department.name if e.department else "Institution-wide",
             e.event_date, e.participants_count]
            for e in events
        ],
    )
    document.add_paragraph()


def _iiqa_dvv_section(document, financial_year):
    document.add_heading("IIQA and DVV Clarification Status", level=1)

    submissions_qs = AccreditationSubmission.objects.select_related("financial_year")
    if financial_year:
        submissions_qs = submissions_qs.filter(financial_year=financial_year)
    submissions = list(submissions_qs.order_by("-financial_year__start_date", "submission_type"))

    if not submissions:
        document.add_paragraph("No IIQA or DVV clarification submissions recorded yet.", style="Intense Quote")
        document.add_paragraph()
        return

    _add_table(
        document,
        ["Type", "Year", "Status", "Submitted At", "Signed Off At"],
        [
            [s.get_submission_type_display(), s.financial_year.label, s.get_status_display(),
             s.submitted_at.date() if s.submitted_at else "—", s.signed_off_at.date() if s.signed_off_at else "—"]
            for s in submissions
        ],
    )
    document.add_paragraph()


def _audited_financials_section(document):
    document.add_heading("5-Year Audited Financial Statements", level=1)

    cert_type = ComplianceCertificateType.objects.filter(name=AUDITED_FINANCIAL_STATEMENT_TYPE_NAME).first()
    trailing_years = list(FinancialYear.objects.order_by("-start_date")[:TRAILING_FINANCIAL_YEARS])

    rows = []
    for fy in trailing_years:
        certificate = None
        if cert_type:
            certificate = ComplianceCertificate.objects.filter(
                certificate_type=cert_type, financial_year=fy,
            ).order_by("-uploaded_at").first()
        rows.append([fy.label, "Yes" if certificate else "No — MISSING"])

    if not rows:
        document.add_paragraph("No financial years have been set up yet.", style="Intense Quote")
    else:
        _add_table(document, ["Financial Year", "Audited Statement on File?"], rows)
    document.add_paragraph()


def build_naac_document(financial_year=None):
    """Returns an io.BytesIO holding the finished .docx (AQAR if financial_year
    is given, otherwise the cumulative SSR)."""
    document = Document()

    _title_page(document, financial_year)
    _extended_profile_section(document)
    _criteria_section(document, financial_year)
    _faculty_output_section(document, financial_year)
    _student_feedback_section(document, financial_year)
    _institutional_events_section(document, financial_year)
    _iiqa_dvv_section(document, financial_year)
    _audited_financials_section(document)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
