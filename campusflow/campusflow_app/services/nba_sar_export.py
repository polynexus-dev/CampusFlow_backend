"""
NBA Self-Assessment Report (SAR) document generator.

Compiles a program's NBA-tagged accreditation criteria (applied narrative +
evidence, from models/accreditation_narrative.py and models/compliance.py)
and its PO/PSO/PEO attainment (services/outcome_attainment.py — the same
computation ProgramOutcomeAttainmentView exposes as CSV) into a single
submission-shaped .docx.

This never generates new "official" content: a criterion with no applied
AccreditationNarrativeDraft renders an explicit "not yet reviewed by IQAC"
placeholder rather than fabricated prose, the same human-in-the-loop
boundary AccreditationNarrativeDraft's own docstring describes.
"""
import io
from datetime import date

from django.db import connection
from django.db.models import Q

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..models.compliance import AccreditationCriterion
from ..models.outcomes import POCOMapping, ProgramOutcome
from ..models.profile import StudentProfile, TeachingStaffProfile
from .outcome_attainment import compute_program_outcome_attainment

NOT_REVIEWED_PLACEHOLDER = (
    "No IQAC-approved narrative has been applied for this criterion/year yet. "
    "Add evidence and request/apply an AI narrative draft before final submission."
)


def _add_table(document, header, rows):
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = "" if value is None else str(value)
    return table


def _title_page(document, program, financial_year):
    tenant_name = getattr(connection.tenant, "name", "") or connection.schema_name

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(tenant_name)
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("NBA Self-Assessment Report (SAR)")
    run.bold = True
    run.font.size = Pt(16)

    program_line = document.add_paragraph()
    program_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    program_line.add_run(f"{program.name} ({program.code}) — {program.get_level_display()}")

    meta_line = document.add_paragraph()
    meta_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fy_label = financial_year.label if financial_year else "All applied narratives (no year filter)"
    meta_line.add_run(f"Financial Year: {fy_label}")

    generated_line = document.add_paragraph()
    generated_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_line.add_run(f"Generated on {date.today().isoformat()} via CampusNexus")

    document.add_page_break()


def _program_particulars_section(document, program):
    document.add_heading("Program Particulars", level=1)

    total_students = StudentProfile.objects.filter(program=program).count()
    total_faculty = TeachingStaffProfile.objects.filter(department=program.department).count()
    ratio = round(total_students / total_faculty, 2) if total_faculty else None
    regulations = list(program.regulations.order_by("-effective_from_year").values_list("code", flat=True))

    _add_table(
        document,
        ["Particular", "Value"],
        [
            ["Department", program.department.name],
            ["Level", program.get_level_display()],
            ["Duration (years)", program.duration_years],
            ["Total Credits Required", program.total_credits_required],
            ["Regulation(s)", ", ".join(regulations) or "—"],
            ["Total Students Enrolled", total_students],
            ["Total Teaching Faculty", total_faculty],
            ["Student : Faculty Ratio", f"{ratio}:1" if ratio else "—"],
        ],
    )
    document.add_paragraph()


def _vision_mission_peo_section(document, program):
    document.add_heading("Vision, Mission and Program Educational Objectives", level=1)

    peos = ProgramOutcome.objects.filter(program=program, kind="peo").order_by("order", "code")
    if peos:
        _add_table(
            document,
            ["Code", "Statement"],
            [[peo.code, peo.statement] for peo in peos],
        )
    else:
        document.add_paragraph(
            "No Program Educational Objectives have been recorded for this program yet."
        )
    document.add_paragraph()


def _criteria_section(document, program, financial_year):
    document.add_heading("Criterion-wise Narrative and Evidence", level=1)

    criteria = AccreditationCriterion.objects.filter(body=AccreditationCriterion.BODY_NBA).order_by("code")
    for criterion in criteria:
        document.add_heading(f"Criterion {criterion.code}: {criterion.title}", level=2)

        narrative_qs = criterion.narrative_drafts.filter(status="applied")
        if financial_year:
            narrative_qs = narrative_qs.filter(financial_year=financial_year)
        narrative = narrative_qs.order_by("-applied_at").first()
        document.add_paragraph(narrative.narrative_text if narrative else NOT_REVIEWED_PLACEHOLDER)

        evidence_qs = criterion.evidence.filter(
            Q(department=program.department) | Q(department__isnull=True)
        )
        if financial_year:
            evidence_qs = evidence_qs.filter(financial_year=financial_year)
        evidence_items = list(evidence_qs.order_by("-created_at"))

        if evidence_items:
            _add_table(
                document,
                ["Description", "Linked Source", "Status"],
                [
                    [
                        item.description or (item.file.name.split("/")[-1] if item.file else "—"),
                        item.linked_object_type or (item.linked_certificate.certificate_type.name if item.linked_certificate else "—"),
                        item.get_status_display(),
                    ]
                    for item in evidence_items
                ],
            )
        else:
            document.add_paragraph("No evidence recorded for this criterion/year yet.", style="Intense Quote")
        document.add_paragraph()


def _co_po_matrix_section(document, program):
    document.add_heading("CO-PO Articulation Matrix", level=1)

    mappings = (
        POCOMapping.objects
        .filter(course_outcome__course__regulation__program=program)
        .select_related("course_outcome__course", "program_outcome")
        .order_by("course_outcome__course__course_code", "course_outcome__code", "program_outcome__code")
    )
    if mappings:
        _add_table(
            document,
            ["Course", "Course Outcome", "Program Outcome", "Correlation Strength"],
            [
                [
                    m.course_outcome.course.course_code,
                    m.course_outcome.code,
                    m.program_outcome.code,
                    m.strength,
                ]
                for m in mappings
            ],
        )
    else:
        document.add_paragraph("No CO-PO mappings have been recorded for this program's courses yet.")
    document.add_paragraph()


def _po_attainment_section(document, program):
    document.add_heading("Program Outcome Attainment", level=1)

    results = compute_program_outcome_attainment(program.id)
    if results:
        rows = []
        for po in results:
            contributing = "; ".join(
                f"{c['course_code']} {c['course_outcome_code']} (strength {c['correlation_strength']}, {c['attainment_value']}%)"
                for c in po["contributing_course_outcomes"]
            )
            rows.append([po["code"], po["kind"].upper(), po["statement"], po["attainment_percent"], contributing])
        _add_table(document, ["Code", "Kind", "Statement", "Attainment %", "Contributing Course Outcomes"], rows)
    else:
        document.add_paragraph("No attainment could be computed — no CO-PO mappings or evaluated exams found yet.")


def build_nba_sar_document(program, financial_year=None):
    """Returns an io.BytesIO holding the finished .docx."""
    document = Document()

    _title_page(document, program, financial_year)
    _program_particulars_section(document, program)
    _vision_mission_peo_section(document, program)
    _criteria_section(document, program, financial_year)
    _co_po_matrix_section(document, program)
    _po_attainment_section(document, program)

    import io
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
